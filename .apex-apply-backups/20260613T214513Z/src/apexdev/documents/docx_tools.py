from __future__ import annotations
from pathlib import Path
import re


def _require_docx():
    try:
        import docx
    except ImportError as e:  # pragma: no cover - optional dependency
        raise ImportError("python-docx is required. Install with `pip install 'apex-dev-tools[docs]'`.") from e
    return docx


def _escape_cell(text: object) -> str:
    return str(text).replace("|", r"\|").replace("\n", "<br>").strip()


def _paragraph_to_markdown(paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    style = getattr(getattr(paragraph, "style", None), "name", "") or ""
    m = re.match(r"Heading\s+(\d+)", style)
    if m:
        level = max(1, min(6, int(m.group(1))))
        return f"{'#' * level} {text}"
    lowered = style.lower()
    if "list bullet" in lowered:
        return f"- {text}"
    if "list number" in lowered:
        return f"1. {text}"
    return text


def docx_to_markdown(path: str | Path, *, include_tables: bool = True) -> str:
    """Convert DOCX paragraphs and tables to Markdown.

    This is still intentionally conservative, but it now preserves heading
    levels, basic list paragraphs, and tables instead of flattening the whole
    document into anonymous paragraphs.
    """
    docx = _require_docx()
    doc = docx.Document(str(path))
    blocks: list[str] = []
    for paragraph in doc.paragraphs:
        line = _paragraph_to_markdown(paragraph)
        if line:
            blocks.append(line)
    if include_tables:
        for table in doc.tables:
            rows = [[_escape_cell(cell.text) for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            width = max(len(r) for r in rows)
            rows = [r + [""] * (width - len(r)) for r in rows]
            header = rows[0]
            blocks.append("| " + " | ".join(header) + " |")
            blocks.append("| " + " | ".join(["---"] * width) + " |")
            for row in rows[1:]:
                blocks.append("| " + " | ".join(row) + " |")
    return "\n\n".join(blocks)


def _add_table(doc, lines: list[str]) -> None:
    rows = []
    for line in lines:
        parts = [c.strip().replace(r"\|", "|") for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", p) for p in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for r_i, row in enumerate(rows):
        for c_i in range(width):
            table.cell(r_i, c_i).text = row[c_i] if c_i < len(row) else ""


def markdown_to_docx(markdown: str, out_path: str | Path) -> Path:
    """Convert a practical Markdown subset to DOCX.

    Supports headings, bullet/numbered lists, blockquotes, fenced-code blocks,
    horizontal rules, simple pipe tables, and paragraphs. For more complete
    Markdown fidelity, use the unified converter with Pandoc installed.
    """
    docx = _require_docx()
    doc = docx.Document()
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            code_lines = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            _add_table(doc, table_lines)
            table_lines = []

    for raw in markdown.splitlines():
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        stripped = line.strip()
        if not stripped:
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
        elif re.match(r"^[-*+]\s+", stripped):
            doc.add_paragraph(re.sub(r"^[-*+]\s+", "", stripped), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
        elif stripped.startswith(">"):
            doc.add_paragraph(stripped.lstrip(">").strip(), style="Intense Quote")
        elif re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped):
            doc.add_paragraph("―" * 24)
        else:
            doc.add_paragraph(stripped)
    flush_table()
    flush_code()
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
